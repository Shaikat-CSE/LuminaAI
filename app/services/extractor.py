"""
Document extraction service for various file types.
Handles PDF, DOCX, TXT, Images (JPG/PNG), CSV, and SQLite files.
"""
import io
import csv
import sqlite3
import base64
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

from app.core.config import get_settings


settings = get_settings()

# Configure Tesseract path for Windows if specified
if settings.tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path


class DocumentExtractor:
    """Extract text from various document types."""
    
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".csv": "csv",
        ".db": "sqlite",
        ".sqlite": "sqlite",
        ".sqlite3": "sqlite",
    }
    
    def __init__(self):
        self.settings = get_settings()
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """Determine file type from extension."""
        ext = Path(filename).suffix.lower()
        return self.SUPPORTED_EXTENSIONS.get(ext)
    
    def extract(self, file_path: Path) -> str:
        """Extract text from a file based on its type."""
        file_type = self.get_file_type(file_path.name)
        
        if not file_type:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_txt,
            "image": self._extract_image,
            "csv": self._extract_csv,
            "sqlite": self._extract_sqlite,
        }
        
        return extractors[file_type](file_path)
    
    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF.
        Uses PyMuPDF first, falls back to pdfplumber.
        Applies OCR only if no selectable text is found.
        """
        text_parts = []
        has_text = False
        
        # Try PyMuPDF first (faster for text-based PDFs)
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    has_text = True
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            doc.close()
        except Exception as e:
            print(f"PyMuPDF failed: {e}, trying pdfplumber...")
        
        # If no text found with PyMuPDF, try pdfplumber
        if not has_text:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            has_text = True
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            except Exception as e:
                print(f"pdfplumber failed: {e}")
        
        # If still no text, apply OCR (scanned PDF)
        if not has_text:
            text_parts = self._ocr_pdf(file_path)
        
        return "\n\n".join(text_parts)
    
    def _ocr_pdf(self, file_path: Path) -> list[str]:
        """Apply OCR to a scanned PDF."""
        text_parts = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                # Render page to image
                pix = page.get_pixmap(dpi=300)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                # Apply OCR
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1} - OCR]\n{page_text}")
            doc.close()
        except Exception as e:
            print(f"OCR failed for PDF: {e}")
        
        return text_parts
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs)
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # Fallback: read as binary and decode with errors ignored
        return file_path.read_bytes().decode('utf-8', errors='ignore')
    
    def _extract_image(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    def extract_image_from_base64(self, base64_string: str) -> str:
        """Extract text from base64-encoded image using OCR."""
        # Remove data URL prefix if present
        if "," in base64_string:
            base64_string = base64_string.split(",")[1]
        
        image_data = base64.b64decode(base64_string)
        image = Image.open(io.BytesIO(image_data))
        text = pytesseract.image_to_string(image)
        return text
    
    def _extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV file."""
        text_parts = []
        
        # Detect encoding
        content = file_path.read_bytes()
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                decoded = content.decode(encoding)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        else:
            decoded = content.decode('utf-8', errors='ignore')
        
        reader = csv.reader(io.StringIO(decoded))
        
        for row_num, row in enumerate(reader):
            if row:
                row_text = " | ".join(str(cell).strip() for cell in row if str(cell).strip())
                if row_text:
                    if row_num == 0:
                        text_parts.append(f"[Headers] {row_text}")
                    else:
                        text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _extract_sqlite(self, file_path: Path) -> str:
        """Extract data from SQLite database."""
        text_parts = []
        
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            
            # Get all table names
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = cursor.fetchall()
            
            for (table_name,) in tables:
                text_parts.append(f"\n[Table: {table_name}]")
                
                # Get column names
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = [col[1] for col in cursor.fetchall()]
                text_parts.append(f"Columns: {', '.join(columns)}")
                
                # Get first 100 rows of data
                cursor.execute(f"SELECT * FROM {table_name} LIMIT 100")
                rows = cursor.fetchall()
                
                for row in rows:
                    row_text = " | ".join(str(cell) for cell in row)
                    text_parts.append(row_text)
            
            conn.close()
        except Exception as e:
            print(f"SQLite extraction error: {e}")
            raise ValueError(f"Could not read SQLite database: {e}")
        
        return "\n".join(text_parts)


# Singleton instance
document_extractor = DocumentExtractor()
